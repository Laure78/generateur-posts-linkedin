#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Génère un compte rendu de chantier 3D MANAGER (.docx) à la charte graphique.

Usage :
    python3 generate_cr.py cr_data.json [chemin/sortie.docx]

Le fichier JSON contient les données structurées du CR (voir schéma dans SKILL.md).
Le rendu (couleurs, logo, polices) est piloté par le bloc CHARTE ci-dessous :
il suffit de remplacer les valeurs par la charte réelle de 3D MANAGER.
"""

import json
import os
import sys
import tempfile

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ───────────────────────────────────────────────────────────────────────────
# CHARTE GRAPHIQUE 3D MANAGER — officielle (anthracite · rouge · fonds clairs)
# Logo : assets/logo_3dmanager.png (détouré, fond transparent, lisible sur bandeau #2A2A2A)
# ───────────────────────────────────────────────────────────────────────────
PRIMAIRE   = "2A2A2A"   # anthracite — bandeaux foncés, texte fort
ROUGE      = "CC2A2A"   # rouge de marque (le « D ») — accents, filets, statut En attente
SECONDAIRE = "2A2A2A"   # texte corps
CLAIR      = "F2F2F2"   # fonds clairs bandeaux d'info / en-têtes tableau
GRIS_LIGNE = "D9D9D9"   # filets de tableau
GRIS_TXT   = "6B6B6B"   # mentions légales / secondaire
BLANC      = "FFFFFF"
BLEU_CHARTE = "2B5C8A"  # statut Nouveau

POLICE_TITRE = "Calibri"
POLICE_CORPS = "Calibri"

STATUT_COULEUR = {
    "Levé":       "2E7D32",
    "En cours":   "B26A00",
    "En attente": ROUGE,
    "Nouveau":    BLEU_CHARTE,
}

_SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
_SKILL_DIR = os.path.abspath(os.path.join(_SCRIPT_DIR, ".."))
_DEFAULT_LOGO = os.path.join(_SKILL_DIR, "assets", "logo_3dmanager.png")

SOCIETE = "3D MANAGER"
ACCROCHE = "Un facilitateur dans l'acte de construire"
MENTIONS_PIED = (
    "3D MANAGER · Bureau d'études TCE · Maîtrise d'œuvre d'exécution (MOEX) · Certifié ISO 9001"
)
RESEAU_AGENCES = "Siège · Atlantique · Aquitaine · Méditerranée · Île-de-France"
# Coordonnées en repli — à confirmer avec la société
CONTACT_PIED = "3dmanager@3dmanager.fr · www.3dmanager.fr"
# ───────────────────────────────────────────────────────────────────────────


def C(hexstr):
    return RGBColor.from_string(hexstr)


def resolve_logo_path():
    """Cherche le logo sur disque (dev local, Docker Railway, cwd skill)."""
    env = os.environ.get("BEWORK_3DM_LOGO_PATH", "").strip()
    candidates = [
        env,
        _DEFAULT_LOGO,
        os.path.join(os.getcwd(), "assets", "logo_3dmanager.png"),
        os.path.join(os.getcwd(), "skills", "3dmanager-cr-chantier", "assets", "logo_3dmanager.png"),
    ]
    for path in candidates:
        if path and os.path.isfile(path):
            return os.path.abspath(path)
    return None


def prepare_logo_for_dark_banner(src_path):
    """
    Prépare une version lisible sur bandeau anthracite (#2A2A2A) :
    fond blanc/noir → transparent, traits sombres → blanc, rouge de marque conservé.
    """
    try:
        from PIL import Image
    except ImportError:
        return src_path

    im = Image.open(src_path).convert("RGBA")
    out_pixels = []
    for r, g, b, a in im.getdata():
        if a < 25:
            out_pixels.append((0, 0, 0, 0))
            continue
        lum = (r + g + b) / 3
        # Rouge « D » de marque
        if r > 160 and g < 100 and b < 100 and lum < 200:
            out_pixels.append((int(ROUGE[:2], 16), int(ROUGE[2:4], 16), int(ROUGE[4:6], 16), a))
            continue
        # Fond clair ou noir plein → transparent
        if lum > 235 or (lum < 35 and a > 200):
            out_pixels.append((0, 0, 0, 0))
            continue
        # Traits clairs (logo blanc)
        if lum > 170:
            out_pixels.append((255, 255, 255, a))
            continue
        # Traits sombres → blanc pour contraste sur anthracite
        out_pixels.append((255, 255, 255, min(255, a + 40)))

    out = Image.new("RGBA", im.size)
    out.putdata(out_pixels)
    tmp = tempfile.NamedTemporaryFile(suffix="_3dm_logo.png", delete=False)
    out.save(tmp.name, "PNG")
    tmp.close()
    return tmp.name


_logo_temp_path = None
_prepared_logo_cache = None


def logo_path_for_docx():
    global _logo_temp_path, _prepared_logo_cache
    if _prepared_logo_cache:
        return _prepared_logo_cache
    src = resolve_logo_path()
    if not src:
        return None
    prepared = prepare_logo_for_dark_banner(src)
    _prepared_logo_cache = prepared
    if prepared != src:
        _logo_temp_path = prepared
    return prepared


def add_logo_to_paragraph(paragraph, height_cm=2.0):
    path = logo_path_for_docx()
    if not path:
        return False
    run = paragraph.add_run()
    run.add_picture(path, height=Cm(height_cm))
    return True


def add_logo_fallback_text(paragraph):
    r = paragraph.add_run(SOCIETE + "\n")
    style_run(r, 14, True, BLANC, POLICE_TITRE)
    r2 = paragraph.add_run(ACCROCHE)
    style_run(r2, 7, False, BLANC, POLICE_CORPS)


def build_brand_banner_table(doc, data, logo_height_cm=2.0):
    """Bandeau logoté visible dans le corps du document (1re page)."""
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.autofit = False
    no_table_borders(tbl)
    tblPr = tbl._tbl.tblPr
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(int(Cm(17).twips)))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)

    left, right = tbl.rows[0].cells
    left.width = Cm(6.5)
    right.width = Cm(10.5)
    set_cell_bg(left, PRIMAIRE)
    set_cell_bg(right, PRIMAIRE)
    set_cell_margins(left, top=140, bottom=140, left=180, right=100)
    set_cell_margins(right, top=140, bottom=140, left=100, right=200)

    lp = left.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if not add_logo_to_paragraph(lp, logo_height_cm):
        add_logo_fallback_text(lp)

    rp = right.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r1 = rp.add_run("COMPTE RENDU DE CHANTIER")
    style_run(r1, 15, True, BLANC, POLICE_TITRE)
    rp2 = right.add_paragraph()
    rp2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r2 = rp2.add_run(f"N° {data.get('cr_numero', '')}")
    style_run(r2, 11, True, ROUGE, POLICE_TITRE)
    if data.get("date_visite"):
        rp3 = right.add_paragraph()
        rp3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        style_run(rp3.add_run(str(data["date_visite"])[:48]), 8, False, BLANC, POLICE_CORPS)

    add_para(doc, space_after=8)


def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_margins(cell, top=60, bottom=60, left=100, right=100):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for tag, val in (("top", top), ("bottom", bottom), ("start", left), ("end", right)):
        node = OxmlElement(f"w:{tag}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        m.append(node)
    tcPr.append(m)


def no_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "none")
        borders.append(e)
    tblPr.append(borders)


def thin_row_bottom_border(cell, hex_color=GRIS_LIGNE):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    b = OxmlElement("w:bottom")
    b.set(qn("w:val"), "single")
    b.set(qn("w:sz"), "4")
    b.set(qn("w:color"), hex_color)
    borders.append(b)
    tcPr.append(borders)


def style_run(run, size=10, bold=False, color=SECONDAIRE, font=POLICE_CORPS):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = C(color)
    run.font.name = font


def add_para(doc, text="", size=10, bold=False, color=SECONDAIRE,
             align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4, space_before=0, font=POLICE_CORPS):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if text:
        r = p.add_run(text)
        style_run(r, size, bold, color, font)
    return p


def add_page_number_field(paragraph):
    run = paragraph.add_run()
    fldb = OxmlElement("w:fldChar"); fldb.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
    flde = OxmlElement("w:fldChar"); flde.set(qn("w:fldCharType"), "end")
    run._r.append(fldb); run._r.append(instr); run._r.append(flde)
    style_run(run, 8, False, GRIS_TXT)


def build_header(doc, data):
    """Bandeau d'en-tête Word (pages suivantes) — même logo que le corps."""
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    for p in list(header.paragraphs):
        p._element.getparent().remove(p._element)

    tbl = header.add_table(rows=1, cols=2, width=Cm(17))
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.autofit = False
    no_table_borders(tbl)
    tblPr = tbl._tbl.tblPr
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(int(Cm(17).twips)))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)
    left, right = tbl.rows[0].cells
    left.width = Cm(6.5)
    right.width = Cm(10.5)
    set_cell_bg(left, PRIMAIRE)
    set_cell_bg(right, PRIMAIRE)
    set_cell_margins(left, top=100, bottom=100, left=140, right=80)
    set_cell_margins(right, top=100, bottom=100, left=80, right=160)

    lp = left.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if not add_logo_to_paragraph(lp, 1.35):
        add_logo_fallback_text(lp)

    rp = right.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r1 = rp.add_run("COMPTE RENDU DE CHANTIER")
    style_run(r1, 13, True, BLANC, POLICE_TITRE)
    rp2 = right.add_paragraph()
    rp2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r2 = rp2.add_run(f"N° {data.get('cr_numero', '')}")
    style_run(r2, 10, True, ROUGE, POLICE_TITRE)


def build_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    for p in list(footer.paragraphs):
        p._element.getparent().remove(p._element)

    # filet
    pf = footer.add_paragraph()
    pPr = pf._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single"); top.set(qn("w:sz"), "8")
    top.set(qn("w:space"), "1"); top.set(qn("w:color"), ROUGE)
    pbdr.append(top); pPr.append(pbdr)
    pf.paragraph_format.space_after = Pt(2)

    p1 = footer.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(p1.add_run(MENTIONS_PIED), 7, True, PRIMAIRE)
    p2 = footer.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(p2.add_run(RESEAU_AGENCES), 7, False, GRIS_TXT)
    p3 = footer.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(p3.add_run(CONTACT_PIED), 7, False, GRIS_TXT)
    pg = footer.add_paragraph()
    pg.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pg.paragraph_format.space_before = Pt(2)
    rr = pg.add_run("Page ")
    style_run(rr, 8, False, GRIS_TXT)
    add_page_number_field(pg)


def info_block(doc, data):
    """Bloc d'identification de l'opération + tableau présents/absents."""
    rows = [
        ("Opération", data.get("operation", "")),
        ("Adresse", data.get("adresse", "")),
        ("Maîtrise d'œuvre d'exécution", data.get("moex", SOCIETE)),
        ("Architecte", data.get("architecte", "")),
        ("Bureau de contrôle / CSPS", " · ".join(filter(None, [data.get("bureau_controle", ""), data.get("csps", "")]))),
        ("Date de visite", data.get("date_visite", "")),
        ("Rédacteur", data.get("redacteur", "")),
        ("Diffusion", data.get("diffusion", "")),
        ("Météo", data.get("meteo", "")),
    ]
    rows = [(k, v) for k, v in rows if v]
    t = doc.add_table(rows=len(rows), cols=2)
    no_table_borders(t)
    t.columns[0].width = Cm(6)
    t.columns[1].width = Cm(13)
    for i, (k, v) in enumerate(rows):
        c0, c1 = t.rows[i].cells
        c0.width = Cm(6); c1.width = Cm(13)
        set_cell_bg(c0, CLAIR)
        set_cell_margins(c0); set_cell_margins(c1)
        style_run(c0.paragraphs[0].add_run(k), 9, True, PRIMAIRE)
        style_run(c1.paragraphs[0].add_run(v), 9, False, SECONDAIRE)
    add_para(doc, space_after=2)

    presents = data.get("presents", [])
    absents = data.get("absents", [])
    if presents or absents:
        pa = doc.add_table(rows=2, cols=2)
        no_table_borders(pa)
        h0, h1 = pa.rows[0].cells
        set_cell_bg(h0, PRIMAIRE); set_cell_bg(h1, PRIMAIRE)
        set_cell_margins(h0); set_cell_margins(h1)
        style_run(h0.paragraphs[0].add_run("Présents"), 9, True, "FFFFFF")
        style_run(h1.paragraphs[0].add_run("Absents / excusés"), 9, True, "FFFFFF")
        b0, b1 = pa.rows[1].cells
        set_cell_margins(b0); set_cell_margins(b1)
        style_run(b0.paragraphs[0].add_run(" · ".join(presents) if presents else "—"), 9, False, SECONDAIRE)
        style_run(b1.paragraphs[0].add_run(" · ".join(absents) if absents else "—"), 9, False, SECONDAIRE)
    add_para(doc, space_after=6)


def section_title(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single"); left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6"); left.set(qn("w:color"), ROUGE)
    pbdr.append(left); pPr.append(pbdr)
    style_run(p.add_run(text), 12, True, PRIMAIRE, POLICE_TITRE)


def lot_table(doc, lot):
    nom = lot.get("nom", "")
    ent = lot.get("entreprise", "")
    titre = nom + (f"  ·  {ent}" if ent else "")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    style_run(p.add_run(titre), 10.5, True, SECONDAIRE, POLICE_TITRE)

    obs = lot.get("observations", [])
    if not obs:
        return
    t = doc.add_table(rows=1 + len(obs), cols=4)
    no_table_borders(t)
    widths = [Cm(1.8), Cm(11.5), Cm(2.6), Cm(2.8)]
    headers = ["N°", "Observation", "Échéance", "Statut"]
    hdr = t.rows[0].cells
    for j, h in enumerate(headers):
        hdr[j].width = widths[j]
        set_cell_bg(hdr[j], CLAIR)
        set_cell_margins(hdr[j])
        thin_row_bottom_border(hdr[j], ROUGE)
        style_run(hdr[j].paragraphs[0].add_run(h), 8.5, True, PRIMAIRE)
    for i, o in enumerate(obs, start=1):
        cells = t.rows[i].cells
        vals = [o.get("num", ""), o.get("texte", ""), o.get("echeance", ""), o.get("statut", "")]
        for j, v in enumerate(vals):
            cells[j].width = widths[j]
            set_cell_margins(cells[j])
            thin_row_bottom_border(cells[j])
            run = cells[j].paragraphs[0].add_run(v)
            if j == 0:
                style_run(run, 8.5, True, SECONDAIRE)
            elif j == 3:
                col = STATUT_COULEUR.get(v.strip(), GRIS_TXT)
                style_run(run, 8.5, True, col)
            else:
                style_run(run, 8.5, False, SECONDAIRE)


def bullet_list(doc, items, marker="⦿"):
    for it in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(0.3)
        style_run(p.add_run(f"{marker} "), 9.5, True, ROUGE)
        style_run(p.add_run(it), 9.5, False, SECONDAIRE)


def build(data, out_path):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.6)
    sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.0)
    sec.right_margin = Cm(2.0)

    # styles par défaut
    normal = doc.styles["Normal"]
    normal.font.name = POLICE_CORPS
    normal.font.size = Pt(10)
    normal.font.color.rgb = C(SECONDAIRE)

    build_header(doc, data)
    build_footer(doc)

    build_brand_banner_table(doc, data, logo_height_cm=2.15)

    info_block(doc, data)

    if data.get("avancement"):
        section_title(doc, "Avancement général")
        add_para(doc, data["avancement"], 10, False, SECONDAIRE, space_after=4)

    section_title(doc, "Observations par corps d'état")
    for lot in data.get("lots", []):
        lot_table(doc, lot)

    if data.get("points_soldes"):
        section_title(doc, "Points soldés depuis le CR précédent")
        add_para(doc, " · ".join(data["points_soldes"]) +
                 f"   —   {len(data['points_soldes'])} point(s) levé(s).",
                 10, False, "2E7D32", space_after=4)

    if data.get("points_attente"):
        section_title(doc, "Points restant en attente")
        bullet_list(doc, data["points_attente"], marker="•")

    if data.get("decisions"):
        section_title(doc, "Décisions et demandes")
        bullet_list(doc, data["decisions"], marker="⦿")

    pr = data.get("prochaine_reunion")
    if pr:
        section_title(doc, "Prochaine réunion")
        add_para(doc, pr.get("date", ""), 10, True, SECONDAIRE, space_after=2)
        if pr.get("ordre_du_jour"):
            add_para(doc, "Ordre du jour : " + pr["ordre_du_jour"], 10, False, SECONDAIRE, space_after=4)

    mention = data.get("mention_approbation",
                       "Le présent compte rendu est réputé approuvé sans observation écrite "
                       "des destinataires dans un délai de 8 jours.")
    p = add_para(doc, mention, 8.5, False, GRIS_TXT, space_before=8)
    p.runs[0].italic = True

    doc.save(out_path)
    return out_path


def main():
    if len(sys.argv) < 2:
        print("Usage : python3 generate_cr.py cr_data.json [sortie.docx]")
        sys.exit(1)
    json_path = sys.argv[1]
    with open(json_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    if len(sys.argv) >= 3:
        out = sys.argv[2]
    else:
        num = str(data.get("cr_numero", "")).zfill(2)
        op = "".join(c for c in data.get("operation", "CR")[:25] if c.isalnum() or c in " -_").strip().replace(" ", "_")
        out = f"CR_{op}_N{num}.docx"

    path = build(data, out)
    if resolve_logo_path():
        print(f"✓ CR généré : {path} (logo 3D MANAGER intégré)")
    else:
        print(f"✓ CR généré : {path} (ATTENTION : logo introuvable — texte de repli)")
    if _logo_temp_path and os.path.isfile(_logo_temp_path):
        try:
            os.unlink(_logo_temp_path)
        except OSError:
            pass


if __name__ == "__main__":
    main()
